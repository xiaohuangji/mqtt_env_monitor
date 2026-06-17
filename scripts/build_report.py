# -*- coding: utf-8 -*-
"""生成《基于 MQTT 的多节点环境监测系统设计》课程设计报告 Word 文档。
按课程模板格式：封面 + 目录 + 8 节正文 + 参考文献；节标题四号黑体、小节小四黑体、
正文小四宋体首行缩进 2 字 1.25 倍行距；图表编号；GB/T 参考文献。
运行：python scripts/build_report.py
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.getcwd()
IMG = "docs/project_outputs/课程设计报告/images"
EXP = "docs/project_outputs/测试证据"
PPTIMG = "docs/project_outputs/汇报PPT/期末答辩/images"
MIDIMG = "docs/project_outputs/汇报PPT/中期检查/images"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
sec.top_margin = sec.bottom_margin = Cm(2.5)
sec.left_margin = sec.right_margin = Cm(2.6)


def set_font(run, cn="宋体", en="Times New Roman", size=12, bold=False, color=(0, 0, 0)):
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    rf.set(qn("w:eastAsia"), cn)
    rf.set(qn("w:ascii"), en)
    rf.set(qn("w:hAnsi"), en)


def body(text, indent=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    set_font(p.add_run(text), "宋体", "Times New Roman", size)
    return p


def h1(text):
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Pt(0)
    set_font(p.add_run(text), "黑体", "Arial", 14, bold=True)
    return p


def h2(text):
    p = doc.add_paragraph(style="Heading 2")
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(text), "黑体", "Arial", 12, bold=True)
    return p


def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), "宋体", "Times New Roman", 10.5)


def figure(path, cap, width_cm=13.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(width_cm))
    else:
        set_font(p.add_run("[图片缺失: %s]" % path), color=(200, 0, 0))
    caption(cap)


def bullets(items, ordered=False):
    style = "List Number" if ordered else "List Bullet"
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(1)
        set_font(p.add_run(it), "宋体", "Times New Roman", 12)


def table(headers, rows, cap=None, widths=None):
    if cap:
        caption(cap)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(c.paragraphs[0].add_run(h), "黑体", "Arial", 10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].paragraphs[0].line_spacing = None
            set_font(cells[i].paragraphs[0].add_run(str(v)), "宋体", "Times New Roman", 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


# ============ 封面 ============
for _ in range(3):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("《数据通信与网络技术》"), "黑体", "Arial", 22, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("课程设计报告"), "黑体", "Arial", 22, bold=True)
for _ in range(4):
    doc.add_paragraph()
cover = [
    ("课设题目", "基于 MQTT 的多节点环境监测系统设计"),
    ("所在小组", "第 ＿＿ 组"),
    ("小组成员", "胡艺瀚　孙博宇　韦煜城　曾家豪"),
    ("指导老师", "和洁　傅攀峰　刘若辰"),
    ("时　　间", "2026 年 6 月"),
]
for k, v in cover:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    set_font(p.add_run("%s：%s" % (k, v)), "宋体", "Times New Roman", 14)
for _ in range(2):
    doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("信息工程学院"), "宋体", "Times New Roman", 14)
doc.add_page_break()

# ============ 目录 ============
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
set_font(p.add_run("目　录"), "黑体", "Arial", 18, bold=True)
tp = doc.add_paragraph()
run = tp.add_run()
f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = 'TOC \\o "1-2" \\h \\z \\u'
f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "separate")
tt = OxmlElement("w:t"); tt.text = "（在 WPS / Word 中右键“更新域”或按 F9 生成目录与页码）"
f3 = OxmlElement("w:fldChar"); f3.set(qn("w:fldCharType"), "end")
for e in (f1, it, f2, tt, f3):
    run._element.append(e)
set_font(run, "宋体", "Times New Roman", 11)
doc.add_page_break()

# ============ 1 设计背景与任务分析 ============
h1("1  设计背景与任务分析")
h2("1.1  设计背景")
body("随着物联网（IoT）与智慧校园、智能农业、机房监测等应用的发展，对温度、湿度、光照、空气质量等环境参数进行多点、实时、远程的采集与汇聚已成为常见需求。这类场景的共同特点是：监测节点数量多、分布广、单点数据量小、上传频率较低，且要求服务端能同时接入大量节点并实时展示。")
body("MQTT（Message Queuing Telemetry Transport）是一种基于发布/订阅模型的轻量级消息传输协议，运行于 TCP 之上，具有报文头开销小、一次连接长期保持、支持三种服务质量（QoS）等级、支持遗嘱与会话保持等特点，非常适合多节点、低带宽、发布/订阅式的物联网通信场景。本课程设计即围绕 MQTT 协议，设计并实现一个多节点环境监测系统，以理解物联网通信中的主题、发布、订阅与消息质量等级等核心机制。")
h2("1.2  设计任务")
body("本题（题目 12）要求：设计并实现一个基于 MQTT 的多节点环境监测系统，多个采集节点通过 MQTT 协议发布环境数据，监控端订阅相关主题并显示数据。基本功能要求包括：至少包含 2 个环境监测节点（可用 ESP32、树莓派或 PC 模拟）；每个节点采集或模拟一种环境数据；搭建 MQTT Broker；设计合理的主题结构；监控端订阅主题并显示节点编号、数据类型、数据值和上传时间。")
body("工程化设计要求进一步规定：分析多节点监测系统的需求与可扩展性；设计 MQTT 主题命名规则；比较 MQTT 与 HTTP 在多节点上传场景中的适用性；分析 QoS 0/1/2 的特点与适用场景；考虑节点掉线、重复消息、消息延迟与服务器负载等问题；并从可扩展性、可靠性、功耗与维护性等方面分析约束。测试与验证要求包括多节点并发、不同 QoS 等级、断线重连测试，并使用 MQTTX、Wireshark 或服务器日志统计消息到达率、平均延迟与异常情况。")
h2("1.3  报告主线与主要工作")
body("中期检查时，指导老师特别强调“方案应由需求驱动，并充分考虑成本，避免一味追求‘最强’”。据此，我们将课题深化为一个更具工程意义的问题：在满足采集需求与服务质量（到达率、延迟、重连）的前提下，网络层服务器所需的最小资源——也就是最低成本——是多少？整份报告围绕这一主线展开。")
body("围绕该主线，本小组完成的主要工作包括：①以 PC 模拟节点与 Web 监控端打通完整通信链路；②MQTT 方案升级，支持 MQTT 5.0（Topic Alias、会话过期）与 json / json-min / msgpack 三种消息格式，正式 Broker 选用 EMQX；③建立“需求参数→网络负载→服务器最小配置→成本”的通用计算方法，并实现需求-成本计算器；④通过抓包、容量/吞吐压测、QoS×丢包、跨公网延迟等实验标定模型系数；⑤接入 ESP32 硬件节点并完成约 15 小时的室外长时间真机验证。")

# ============ 2 需求分析与设计指标 ============
h1("2  需求分析与设计指标")
h2("2.1  应用场景与功能需求")
body("本系统面向智慧校园的全域环境监测这一具体应用。校园中绿化带、户外公共区域、实验楼等分布广泛，温湿度、光照、土壤墒情、空气质量等环境参数目前多依靠人工抽检，存在不实时、覆盖不全、人力成本高等问题；绿化灌溉缺乏数据支撑，实验楼、食堂等还存在燃气泄漏等安全隐患。因此需要一套覆盖全校、实时上报、可扩展的多节点环境监测系统，为绿化按需灌溉、环境舒适度分析与烟雾/燃气异常告警提供数据支撑。所采集的五类数据各有明确用途：温湿度用于微气候与舒适度评估，光照用于植物生长与景观照明，土壤湿度用于绿化灌溉，可燃气体用于空气质量与燃气安全预警。")
body("从节点规模看，该场景跨度很大：单个实验室或楼层为十几到几十个节点，单栋楼或院系为上百个节点，而校园全域可达万级节点。同一套方案需要能够从几十个节点平滑扩展到万级，这是需求分析中最关键的工程特征。")
body("据此，系统的功能需求归纳为：①多节点采集或模拟环境数据，上传周期可配置；②各节点采用统一的主题与消息格式，经 Broker 转发；③Web 监控端订阅主题并实时显示各节点的编号、数据类型、数值与上传/接收时间；④节点断线后能自动重连，离线期间的消息不丢失；⑤对读取失败等异常数据进行识别与提示；⑥全过程留存日志等测试证据。")
h2("2.2  性能与服务质量指标")
body("为使需求“可衡量、可验证”，本系统定义了如下硬性服务质量（QoS）指标，作为后续设计与测试的验收基准。这些指标独立于节点规模与服务器配置，不随规模扩展或成本优化而放宽，是方案设计与配置选型必须守住的底线。")
table(["指标", "目标值", "说明"],
      [["消息到达率", "≥ 95 %", "按消息序号 seq 逐条统计，实收/应收"],
       ["平均传输延迟", "≤ 1 s", "节点发布到监控端接收的时间差"],
       ["断线重连恢复", "≤ 10 s", "网络中断后自动恢复并续传的时间"],
       ["连续运行稳定性", "≥ 30 min 无异常", "长时间运行不崩溃、不掉线"],
       ["并发节点数", "≥ 2（并验证万级）", "多节点同时上传"]],
      cap="表 2-1  性能与服务质量指标")
h2("2.3  工程约束与“最低成本”目标")
body("结合中期反馈，本系统的设计目标被重新定义为一个有约束的最优化问题：不追求最大并发或最强性能，而是在覆盖采集需求、且各项 QoS 指标达标的前提下，求得最低成本的方案。为使问题边界清晰，本阶段将“成本”限定在网络层服务器资源，即带宽、CPU、内存与存储；传感器硬件与节点供电成本不纳入本阶段范围，留待硬件阶段扩展。这一约束把抽象的“做系统”转化为可计算的“求一个有约束的最小值”，并直接引出第 5 章的网络负载与成本建模、以及第 6 章的实验校准。")

# ============ 3 总体方案设计 ============
h1("3  总体方案设计")
h2("3.1  系统总体架构")
body("系统采用“感知层—网络层—应用层”的三层结构，如图 3-1 所示。感知层由 ESP32 硬件节点与 PC 模拟节点构成，负责采集或模拟环境数据并以 MQTT 协议发布；网络层以 EMQX 作为 MQTT Broker，经 Wi-Fi 或以太网接入，完成消息的统一接收与转发；应用层包括 Web 监控端（实时显示）与需求-成本计算器（容量与成本估算）。")
figure(os.path.join(IMG, "系统架构图.png"), "图 3-1  系统总体架构图", 14.5)
h2("3.2  模块划分")
bullets([
    "采集/模拟节点：ESP32 硬件节点采集温/湿/光/可燃气体/土壤等数据；PC 模拟节点以均值回归随机游走生成平滑数据，用于在硬件就绪前打通链路与规模化压测。",
    "MQTT Broker（EMQX）：负责连接管理、主题匹配与消息转发，是系统的网络层核心。",
    "Web 监控端：基于 Flask + SSE，订阅全部主题，实时显示并记录毫秒级接收日志。",
    "需求-成本计算器：将网络负载与服务器配置的建模方法固化为工具，按需求参数输出最小配置与成本。",
])
h2("3.3  数据流向与接口")
body("数据流向为：节点按主题 env_monitor/{node_id}/{data_type} 发布（publish）→ EMQX Broker 接收并按订阅关系转发 → Web 监控端以通配符 env_monitor/+/+ 订阅（subscribe）并显示。系统的关键设计在于节点接口统一：ESP32 硬件节点与 PC 模拟节点采用完全相同的主题与消息格式接入，监控端无需区分数据来源。正是这一统一接口，使得“先以软件节点打通链路、再接入硬件节点”的工程路线得以平滑实施。")

# ============ 4 关键技术与协议设计 ============
h1("4  关键技术与协议设计")
h2("4.1  通信协议选择：MQTT 与 HTTP 的比较")
body("在多节点数据上传场景中，本系统在 MQTT 与 HTTP 之间进行了比较。HTTP 采用请求/响应模型，客户端需主动轮询，且每次请求都要建立连接、携带较大的报文头，在节点多、上传频繁时开销和时延都较大，也不便于服务端主动推送。MQTT 采用发布/订阅模型，一次连接长期保持，报文头极小（最小固定头仅 2 字节），支持服务端向订阅者主动推送、支持 QoS 分级与遗嘱机制，天然契合多节点、低带宽、实时汇聚的物联网场景。因此本系统选用 MQTT 作为应用层协议。相比直接使用 TCP/UDP 裸 Socket，MQTT 还免去了自行实现连接保持、消息路由与可靠传输的工作量。")
h2("4.2  主题结构与消息格式设计")
body("主题命名采用三段式：env_monitor/{node_id}/{data_type}，例如 env_monitor/node01/temperature。其中 node_id 仅标识节点身份，不与数据类型绑定，一个节点可发布多种数据。该结构便于按需订阅：env_monitor/+/+ 订阅全部数据，env_monitor/node01/+ 按节点订阅，env_monitor/+/temperature 按数据类型订阅。")
body("消息载荷统一包含节点编号、数据类型、数值、单位、时间戳与序号（seq）等字段，其中 seq 为测试统计预留，用于逐条核对到达率。为研究带宽与成本，本系统实现了三种等价的消息格式：完整 json、压缩短键的 json-min 以及二进制 msgpack，其字段映射与单条实测字节数如表 4-1 所示（含 TCP/IP 头、QoS 1 交互与 Topic Alias）。")
table(["格式", "字段写法", "单条实测开销", "说明"],
      [["json", "node_id/data_type/value/unit/timestamp/seq", "261.9 B", "完整可读，便于调试"],
       ["json-min", "短键 n/d/v/t/s，单位由监控端反查", "203.0 B", "兼顾可读与紧凑"],
       ["msgpack", "二进制编码短键", "185.9 B", "最省，比 json 小约 29 %"]],
      cap="表 4-1  三种消息格式的字段与单条开销对比")
h2("4.3  QoS 与可靠性机制")
body("MQTT 提供三种 QoS：QoS 0（至多一次，不确认）、QoS 1（至少一次，带 PUBACK 确认，可能重复）、QoS 2（恰好一次，四次握手，开销最大）。本系统上报选用 QoS 1：它在保证消息不丢失的同时开销适中，配合“持久会话”（clean_session=False + 固定 client_id）可在节点离线期间由 Broker 暂存并在重连后补发。节点设置 keepalive 心跳并启用自动重连，从而应对掉线、重复消息与消息延迟等问题；监控端按 seq 去重与统计，避免重复消息影响结果。")
h2("4.4  MQTT 5.0 与 3.1.1 的选择")
body("MQTT 3.1.1 是目前最普及、最稳定的版本，几乎所有客户端库（含 ESP32 固件）都支持；MQTT 5.0 在其基础上新增了 Topic Alias（主题别名）、会话过期间隔、原因码、用户属性、共享订阅等面向大规模与可观测的特性。本系统中，监控端与 PC 模拟节点使用 MQTT 5.0，借助 Topic Alias 把长主题在连接内替换为短整数别名，从而每条消息节省约 21.3 字节；ESP32 节点使用 MQTT 3.1.1，EMQX 同时支持两个版本，二者混用互不影响。值得说明的是，实测表明 5.0 与 3.1.1 的“裸协议开销”基本持平（约 283 B 与 280 B），5.0 的省流量收益主要来自 Topic Alias 这一特性，而非协议本身更轻。")
h2("4.5  Broker 选型")
body("Broker 是网络层核心，本系统对比了 EMQX、amqtt 与 Mosquitto。EMQX 基于 Erlang/OTP，是生产级商用 Broker，依托 Erlang 轻量进程与多核能力可单机承载几十万乃至百万连接，并自带 Dashboard、规则引擎与监控；amqtt 是纯 Python（asyncio）实现，轻量易用、基底内存低，适合开发测试与小规模，但受 GIL 限制基本单核、每连接开销大，连接过万后内存与 CPU 都难以支撑；Mosquitto 为 C 实现的轻量 Broker，本系统将其降为离线备用。由于本系统面向万级规模，最终选用 EMQX（5.8.x；不选 6.x 是因其企业版按 license 限制连接数）。具体容量对比数据见第 6 章。")
h2("4.6  网络分层分析")
body("从网络分层看：应用层为 MQTT（发布/订阅、主题、QoS）；传输层为 TCP，Broker 默认监听 1883 端口，借助 TCP 的确认与重传保证字节流可靠；网络层为 IP，涉及子网划分与路由，万级接入还需考虑 Wi-Fi AP 的关联容量；链路层为 Wi-Fi 或以太网；物理层为 ESP32 开发板与各类传感器。这一分层关系也是后续容量与成本建模的基础。")

# ============ 5 实现过程与关键配置 ============
h1("5  实现过程与关键配置")
h2("5.1  PC 模拟节点")
body("PC 模拟节点（src/simulated_nodes/sim_node.py）以均值回归随机游走生成平滑且贴近真实环境的数据，避免数值突跳。节点编号、数据类型、上传周期、QoS、消息格式与协议版本均可通过命令行配置；在 MQTT 5.0 下按连接维护 Topic Alias，断线时 QoS 1 消息在本地排队并于重连后补发。模拟节点既用于在硬件就绪前打通链路，也用于规模化压测（配合压测工具可达数万连接）。")
h2("5.2  ESP32 硬件节点")
body("ESP32 硬件节点连接 DHT11（温湿度）、BH1750（光照）、MQ-2（可燃气体）、FC-28（土壤湿度）等传感器，采集后以 MQTT 3.1.1、json-min 格式按约 30 秒周期发布到 Broker，由充电宝供电以支持室外长时间运行。节点接口与模拟节点完全一致，因而可直接接入既有链路而无需改动监控端。")
h2("5.3  Web 监控端")
body("Web 监控端（src/monitor/web_monitor.py）基于 Flask 提供页面，后端用 paho-mqtt 订阅 env_monitor/+/+，并通过服务器推送事件（SSE）将更新实时推送到浏览器，页面无需刷新即可更新各节点的最新数值与在线状态；超过 15 秒未收到数据的条目自动标记为离线/过期。监控端同时记录毫秒级接收日志（含 seq、QoS、格式、原始载荷），为到达率与延迟统计提供证据；并提供 /calculator 路由承载需求-成本计算器。监控端实时界面见图 5-1。")
figure(os.path.join(MIDIMG, "web_monitor_online_crop.png"), "图 5-1  Web 监控端实时显示界面（双节点并发）", 14.0)
h2("5.4  EMQX 部署与系统调优")
body("EMQX 以 Docker 方式部署（emqx/emqx:5.8.9），对外开放 1883（MQTT）与 18083（Dashboard）端口。为支撑万级并发连接，需对操作系统进行调优，主要包括放开文件句柄上限（ulimit）、扩大本地端口范围等内核参数，相关配置已在压测服务器上持久化验证。演示环境下，开启笔记本移动热点后，ESP32 即可接入同一 Broker。")
h2("5.5  需求-成本计算器")
body("需求-成本计算器是“最低成本”主线的工具化实现，其算法遵循“需求参数→网络负载→满足 QoS 的最小配置→成本”的计算链。给定节点数 N、数据类型数 M、上传周期 T、QoS、消息格式、订阅端数 S 等参数，先由消息率 R = N·M / T 与单条开销 B_msg 推出上/下行带宽与月流量；再据此推算服务器配置：内存采用“空闲下界 + 满负载上界”双模型（满负载按 1.25 余量选型），CPU 由 R·(1+S)/k_core·k_s 估算，磁盘按落库字节与保留天数估算，并给出所需 Wi-Fi AP 数。所有关键系数（单条字节数、每连接内存、每核吞吐等）均由第 6 章的抓包与压测实验标定。计算器以网页形式（/calculator）实现，输入需求即可输出最小配置与成本，可用于任意场景。")

# ============ 6 测试方案与结果分析 ============
h1("6  测试方案与结果分析")
h2("6.1  测试环境与项目")
body("测试在三个台架上进行并辅以一个对照组：①本地抓包台架，用 Wireshark 标定单条消息开销；②云服务器（8 核 32 GB）上的 EMQX 压测，测容量与吞吐上限；③跨公网链路，测真实延迟与到达率；④以纯 Python 的 amqtt 作为对照，验证 Broker 选型。所用工具包括 MQTTX、Wireshark、emqtt_bench、自写的序号化延迟探针与服务器日志。测试项目覆盖功能、连通性、协议、性能、异常与稳定性。")
h2("6.2  功能与连通性测试")
body("双节点并发上传测试中，监控端按 seq 与发送日志逐条核对，结果如表 6-1：到达率 100 %、断线后约 7 秒重连、QoS 1 离线期间排队的 12 条消息在重连后全部补发、零丢失，各项功能指标均达标。")
table(["测试项目", "测试方法", "测试结果", "是否通过"],
      [["多节点并发", "双节点同时上传，seq 逐条对照", "到达率 100 %（48/48）", "通过"],
       ["断线重连", "断开网络后恢复，记录恢复时间", "约 7 s（指标 ≤ 10 s）", "通过"],
       ["QoS 1 补发", "离线期间发布，重连后核对", "离线 12 条全部补发，0 丢失", "通过"],
       ["实时显示", "页面与日志比对", "节点号/类型/值/时间实时刷新", "通过"]],
      cap="表 6-1  功能与连通性测试结果")
h2("6.3  协议开销实验（抓包）")
body("用 Wireshark 抓取不同 QoS、格式、是否启用 Topic Alias 及协议版本下的报文，测得单条消息总开销（含 TCP/IP 头、QoS 交互与 ACK 分摊）：QoS 1 + Topic Alias 稳态下 json / json-min / msgpack 分别为 261.9 / 203.0 / 185.9 字节；相对 QoS 1，QoS 0 约省 59 字节、QoS 2 约增 106 字节；Topic Alias 每条节省约 21.3 字节；空闲心跳约 116 字节/周期。如图 6-1。这些字节系数是带宽与月流量计算的依据。")
figure(os.path.join(EXP, "实验图_协议开销对比.png"), "图 6-1  不同 QoS / 格式 / 协议下的单条消息开销", 13.5)
h2("6.4  容量与吞吐压测")
body("在 8 核 32 GB 服务器上对 EMQX 做连接爬坡与吞吐压测：5 万空闲连接仅占用约 1.05 GB 内存（约 15.75 KB/连接），满负载下每连接内存约 38.76 KB（为空闲的 2.46 倍，故选型须按满负载上界）；处理 5 万条/秒消息仅需约 3.2 个核（每核每秒约 1.9 万条）。作为对照，amqtt 在千级连接即出现单核饱和：其基底内存低、约 1 万连接以内反而比 EMQX 更省，但每连接开销约为 EMQX 的 3.6 倍，过万后被迅速反超，5 万连接达 2.79 GB。综合内存、CPU 与横向扩容能力，面向万级规模应选 EMQX。相关曲线见图 6-2 至图 6-4。")
figure(os.path.join(EXP, "实验图_连接容量对比.png"), "图 6-2  EMQX 与 amqtt 连接容量（内存）对比", 13.5)
figure(os.path.join(EXP, "实验图_吞吐对比.png"), "图 6-3  EMQX 吞吐与 CPU 占用", 13.5)
figure(os.path.join(EXP, "实验图_内存空闲vs满负载.png"), "图 6-4  EMQX 空闲与满负载内存双模型", 13.5)
h2("6.5  服务质量（可靠性）实验")
body("QoS×丢包实验在 0–15 % 的人为丢包下测试三种 QoS 的到达率与延迟：由于 MQTT 运行在 TCP 之上，TCP 的重传机制将丢失的报文补回，三种 QoS 的应用层到达率均保持 100 %，代价是延迟随丢包率上升（图 6-5 中红色虚线为“无 TCP 重传时的理论到达率”对照，可见若无重传，丢包下到达率会显著下降）。跨公网延迟实测平均 44.7 ms、到达率 100 %（图 6-6），断线重连约 5.2 秒，均满足指标。")
figure(os.path.join(EXP, "实验图_QoS丢包到达率与延迟.png"), "图 6-5  不同 QoS 在丢包下的到达率与延迟", 13.5)
figure(os.path.join(EXP, "实验图_公网延迟分布.png"), "图 6-6  跨公网传输延迟分布", 12.5)
h2("6.6  实景压测（30 分钟稳定性）")
body("以 2 万连接、每连接每 15 秒发布一条、10 路订阅全量转发的实景负载，连续运行 30 分钟：TCP 连接数全程恒定（零掉线），订阅端转发吞吐稳定在约 1.33 万条/秒、送达率 99.96 %，EMQX 进程 CPU 均值约 85 %（约 1.1 核）、内存稳定在约 1.5 GB（图 6-7）。该结果与计算器对同等规模的配置预测吻合，说明模型可信。")
figure(os.path.join(EXP, "场景6实景压测30分钟曲线.png"), "图 6-7  2 万连接 30 分钟实景压测曲线", 11.5)
h2("6.7  ESP32 真机长时间测试")
body("2026 年 6 月 15 日，ESP32 硬件节点在室外连续运行约 15 小时（07:19–22:18），由云端记录器持久记录。结果：共上报 8547 条，整体到达率 95.0 %（期间发生 2 次网络中断，扣除后稳态接近 100 %），QoS 1 与持久会话保证了中断恢复后的数据自动续传、零不可恢复丢失。温度、湿度、可燃气体、土壤湿度四类数据合理（见表 6-3 与图 6-8），其中温度日变化（清晨约 28 ℃、午后峰值约 39 ℃、傍晚回落）与当日西安晴热高温天气高度吻合，验证了系统数据的真实性。")
table(["传感器", "数据量", "实测范围", "状态"],
      [["温度（DHT11）", "1711", "26.2 ~ 39.3 ℃", "正常"],
       ["湿度（DHT11）", "1709", "29 ~ 98 %", "正常"],
       ["可燃气体（MQ-2）", "1709", "310 ~ 696（原始值）", "正常"],
       ["土壤湿度（FC-28）", "1709", "66.65 ~ 70.02 %", "正常"]],
      cap="表 6-3  ESP32 室外四类环境数据实测（2026-06-15）")
figure(os.path.join(PPTIMG, "esp32_sensors.png"), "图 6-8  ESP32 室外四类环境数据全天变化（2026-06-15）", 13.5)
body("光照方面，BH1750 经 I2C 接口采集，在 6 月 16 日凌晨至清晨连续约 6 小时稳定有效（夜间 0 lux、日出后平滑爬升至约 305 lux，见图 6-9），当日白天进一步测得约 8000 lux，正确反映了昼夜光照变化。早期白天曾因 ESP32 端接线接触不良出现间歇性读取缺失，重新整理接线、共地并加入数据有效性判断后恢复正常。")
figure(os.path.join(IMG, "光照夜晨曲线.png"), "图 6-9  ESP32 光照实测·夜间至清晨连续采集（2026-06-16）", 13.0)
h2("6.8  案例测算与模型互验")
body("将方法落到一个具体案例：校园全域监测，每 100 m² 布设一个节点。以渭水校区约 200 万 m² 估算，约需 2 万个节点。把该需求（2 万节点、4 类、周期 60 s、QoS 1、json-min、10 个监控端）输入计算器，得到表 6-4 的结果：消息率约 1333 条/s、总流量约 26.4 Mbps、月流量约 8.6 TB、30 天落库约 345 GB，服务器最小配置为 2 核 4 GB。该配置与 6.6 节的 2 万连接实景压测结果（约 1.1 核、1.5 GB）相互印证，模型可信。可行性方面：带宽需求 26.4 Mbps（留余量约 33 Mbps）远低于校园网 100 Mbps，约占 1/3；存储约 11.5 GB/天，30 天约 345 GB，配 ≥400 GB 磁盘即可。也就是说，覆盖全校 2 万个节点，网络层只需一台约 600–1000 元的 2 核 4 GB 迷你主机，瓶颈在带宽与存储而非 Broker，二者均充足。")
table(["指标", "计算器输出", "可行性"],
      [["节点数 N", "约 20000", "—"],
       ["消息率 R", "约 1333 条/s", "—"],
       ["总流量 W", "约 26.4 Mbps", "< 校园网 100 Mbps，约 1/3"],
       ["月流量", "约 8.6 TB", "—"],
       ["磁盘（30 天）", "约 345 GB", "配 ≥400 GB 盘"],
       ["服务器最小配置", "2 核 4 GB", "迷你主机约 ¥600–1000"]],
      cap="表 6-4  校园全域（2 万节点）案例计算结果")

# ============ 7 工程分析与改进方向 ============
h1("7  工程分析与改进方向")
h2("7.1  多维度工程约束分析")
bullets([
    "实时性：平均传输延迟实测约 44.7 ms，远优于 ≤ 1 s 指标；丢包会增加延迟但不丢消息。",
    "可靠性：QoS 1 + 持久会话 + 自动重连，使网络短时中断后自动续传；功能测试与 15 小时真机测试到达率 95 %–100 %。",
    "可扩展性：EMQX 单机实测可达 5 万连接、远未触顶，万级场景仅占用约 1 GB 内存；规模再增可横向集群。",
    "成本：2 万节点的校园场景，网络层仅需一台约 600–1000 元的 2 核 4 GB 迷你主机，瓶颈在带宽与存储且均充足，体现了“需求驱动、最低成本”的设计目标。",
    "可维护性：EMQX 自带 Dashboard 监控，监控端留存毫秒级日志，便于排障与复现；本课程设计全过程在 Git 上留痕。",
])
h2("7.2  存在的不足")
bullets([
    "光照传感器 BH1750 早期白天因接线接触不良出现间歇性缺失，已通过重新接线、共地与数据有效性判断解决（夜间至清晨连续 6 小时有效）；后续可再补一段完整白天光照日变化曲线。",
    "现场供网采用手机热点，稳定性不足，长时间测试中出现 2 次网络中断，是到达率未达 100 % 的主要原因。",
    "下行带宽随订阅端数 S 放大：10 个监控端全量转发时，月流量主要由转发贡献，订阅端较多时带宽成本不可忽视。",
    "存储随保留期线性增长；当前未做数据库持久化与阈值告警；Broker 采用匿名连接，尚未做认证与加密。",
])
h2("7.3  改进方向")
bullets([
    "硬件：重新加固/重焊 BH1750 接线，固件增加读取失败重试与 I2C 总线复位，稳定后择日白天补测光照日变化曲线；现场改用更稳定的接入方式。",
    "带宽：多监控端改用 MQTT 5.0 共享订阅（$share）或聚合网关，避免下行被 S 倍放大。",
    "存储与功能：超过周/月保留期的数据按小时/天降采样后冷存或丢弃；增加时序数据库持久化与超限阈值告警。",
    "安全：为 Broker 启用 TLS 加密与用户名口令/客户端证书认证，提升涉外部署时的安全性。",
])

# ============ 8 总结与参考资料 ============
h1("8  总结与参考资料")
h2("8.1  总结")
body("本课程设计完成了一个基于 MQTT 的多节点环境监测系统：以 PC 模拟节点与 Flask + SSE 监控端打通完整通信链路，升级支持 MQTT 5.0 与三种消息格式并选用 EMQX 作为正式 Broker，建立了“需求→负载→最小配置→成本”的通用计算方法并实现需求-成本计算器，通过抓包与压测实验标定了模型系数，最后接入 ESP32 硬件节点完成约 15 小时室外真机验证。各项性能与服务质量指标（到达率 ≥ 95 %、延迟 ≤ 1 s、重连 ≤ 10 s、长时间稳定、多节点并发）全部达标。")
body("在课程知识应用方面，本设计综合运用了网络分层结构、MQTT 发布/订阅与 QoS 机制、TCP 的可靠传输、抓包分析以及服务器容量与配置估算等知识。更重要的是，围绕“需求驱动、考虑成本”的主线，我们没有一味追求最强性能，而是论证了在满足需求与服务质量前提下网络层服务器的最小资源，形成了完整的“需求—设计—实现—测试—分析”证据链，也体现了对复杂工程问题中约束权衡的工程判断。")
h2("8.2  参考资料")
refs = [
    "[1] OASIS. MQTT Version 5.0[S/OL]. OASIS Standard, 2019.",
    "[2] OASIS. MQTT Version 3.1.1[S/OL]. OASIS Standard, 2014.",
    "[3] EMQ Technologies. EMQX 5.0 产品文档[EB/OL]. https://docs.emqx.com.",
    "[4] Eclipse Foundation. Eclipse Paho MQTT Python Client 文档[EB/OL].",
    "[5] 谢希仁. 计算机网络[M]. 8 版. 北京: 电子工业出版社, 2021.",
    "[6] Banks A, Briggs E, Borgendale K, et al. MQTT: The Standard for IoT Messaging[EB/OL]. OASIS, 2019.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(2)
    set_font(p.add_run(r), "宋体", "Times New Roman", 10.5)

# 页脚页码
footer_p = sec.footer.paragraphs[0]
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer_p.add_run()
fa = OxmlElement("w:fldChar"); fa.set(qn("w:fldCharType"), "begin")
ic = OxmlElement("w:instrText"); ic.set(qn("xml:space"), "preserve"); ic.text = "PAGE"
fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "end")
for e in (fa, ic, fb):
    r._element.append(e)
set_font(r, "Times New Roman", "Times New Roman", 10.5)

out = "docs/project_outputs/课程设计报告/《基于MQTT的多节点环境监测系统》课程设计报告.docx"
doc.save(out)
print("saved:", out, "| 段落数:", len(doc.paragraphs))
